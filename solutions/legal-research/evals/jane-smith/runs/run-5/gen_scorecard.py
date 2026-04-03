"""Generate QA scorecard .docx for run-5."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── helpers ──────────────────────────────────────────────────────────────────
def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def para(text="", bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p

def table_row(tbl, cells, bold_first=False):
    row = tbl.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, cells)):
        cell.text = text
        if bold_first and i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True
    return row

def shade_row(row, hex_color):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

# ── Title ─────────────────────────────────────────────────────────────────────
doc.add_heading("QA Score: Magazine, Inc. v. Jane Smith (2025)", 0)
para(f"Run: run-5  |  Date: 2026-03-31  |  Eval: jane-smith  |  Solution: legal-research", italic=True)
doc.add_paragraph()

# ── Run Log ───────────────────────────────────────────────────────────────────
heading("Run Log", 1)
log_tbl = doc.add_table(rows=1, cols=3)
log_tbl.style = "Table Grid"
hdr = log_tbl.rows[0].cells
hdr[0].text = "Step"
hdr[1].text = "Action"
hdr[2].text = "Observed"
for c in hdr:
    for run in c.paragraphs[0].runs:
        run.bold = True
shade_row(log_tbl.rows[0], "D5E8F0")

log_rows = [
    ("Intake", 'Clicked "Select an example to get started" dropdown, selected "Jane Smith"',
     'Title auto-filled "Magazine, Inc. v. Jane Smith (2025)", Jurisdiction=California, Representing=Defendant, Fact pattern populated (811 chars)'),
    ("Follow-up Questions", "Left all 5 questions at Include default, clicked Continue to Areas of Law",
     "Confirmation modal appeared; clicked Yes to proceed without answering"),
    ("Areas of Law", "Accepted all defaults (3 areas: Telecom/Media/Entertainment, Intellectual Property, Personal Injury), clicked Continue",
     "All 3 areas appeared with correct jurisdiction labels; Litigation Stage defaulted to State Trial Court"),
    ("Legal Authorities", "Clicked each of 3 causes of action in sidebar to mark as reviewed",
     "3/3 reviewed — Copyright Infringement (Federal), Unfair Competition (State-CA), Defamation-Libel (State-CA); all showed green checkmarks"),
    ("Legal Questions", "Clicked into each cause to verify question selection",
     "3/3 selected — all causes auto-selected 2 questions each; no manual selection required"),
    ("Cases", "Clicked each of 3 causes in sidebar to mark as reviewed",
     "3/3 reviewed; cases and holdings populated per cause"),
    ("Mapping & Analysis", "Clicked Continue to Analysis",
     "Holdings mapped (5 for Copyright, 3 for Unfair Competition, 1 for Defamation); advanced to Results & Analysis"),
    ("Results & Analysis", "Waited for Temporary indicator to clear",
     "Analysis summary generated for all 3 causes; per-cause analysis visible; Temporary tag cleared"),
    ("Document Generation", 'Clicked Next Steps → Legal Brief card → Persuasive Analysis (default) → "Generate Brief without uploading a complaint"',
     'Status showed "in_progress" then "Downloading brief..."; .docx downloaded automatically'),
]
for step, action, observed in log_rows:
    r = log_tbl.add_row()
    r.cells[0].text = step
    r.cells[1].text = action
    r.cells[2].text = observed

# set col widths
for row in log_tbl.rows:
    row.cells[0].width = Inches(1.4)
    row.cells[1].width = Inches(2.8)
    row.cells[2].width = Inches(3.2)

doc.add_paragraph()

# ── Score Summary ─────────────────────────────────────────────────────────────
heading("Score Summary", 1)

total_para = doc.add_paragraph()
run = total_para.add_run("Total: 97/100   PASS")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1A, 0x73, 0x1A)

score_tbl = doc.add_table(rows=1, cols=4)
score_tbl.style = "Table Grid"
hdr = score_tbl.rows[0].cells
for i, h in enumerate(["Category", "Score", "Max", "Status"]):
    hdr[i].text = h
    for run in hdr[i].paragraphs[0].runs:
        run.bold = True
shade_row(score_tbl.rows[0], "D5E8F0")

score_data = [
    ("Flow Completion", "10", "10", "PASS"),
    ("Document Structure", "10", "10", "PASS"),
    ("Completeness", "5", "5", "PASS"),
    ("Results Page Coverage", "15", "15", "PASS"),
    ("Legal Argument Coverage", "20", "20", "PASS"),
    ("Case Citations", "10", "10", "PASS"),
    ("Analytical Quality", "27", "30", "PASS"),
]
for cat, score, mx, status in score_data:
    r = score_tbl.add_row()
    r.cells[0].text = cat
    r.cells[1].text = score
    r.cells[2].text = mx
    r.cells[3].text = status
    color = "E8F5E9" if status == "PASS" else "FFEBEE"
    shade_row(r, color)

doc.add_paragraph()

# ── Score Rationale ───────────────────────────────────────────────────────────
heading("Score Rationale", 1)

heading("Structural (25/25)", 2)
rationale_items = [
    ("Flow Completion — 10/10",
     "All 9 steps completed without human intervention. All gates passed (3/3 Legal Authorities, 3/3 Legal Questions, 3/3 Cases). .docx downloaded automatically."),
    ("Document Structure — 10/10",
     "All 12 required sections present in correct order: Legal Advisory & Disclaimer, Court Header (US District Court for the [X] District of California), Table of Authorities, Introduction, Statement of Jurisdiction, Statement of Issues, Statement of the Case, Statement of Facts, Argument, Conclusion, Certificate of Compliance, Certificate of Service."),
    ("Completeness — 5/5",
     "Word count 7,261 (within 5,000–9,000 range per certificate). Parties identified correctly throughout. Conclusion explicitly requests judgment in favor of Ms. Smith."),
]
for title, body in rationale_items:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(title + ": ")
    run.bold = True
    p.add_run(body)

heading("Substantive (72/75)", 2)
sub_items = [
    ("Results Page Coverage — 15/15",
     "All 3 areas of law appeared (Telecom/Media/Entertainment, Intellectual Property, Personal Injury). All 3 required causes present with correct jurisdiction labels: Copyright Infringement (Federal), Unfair Competition (State-CA), Defamation-Libel (State-CA). Cases populated per cause."),
    ("Legal Argument Coverage — 20/20",
     "All 6 required topics substantively argued: (I) Fair Use §107 with full four-factor analysis; (II) Copyright Damages §504(b) with causation/attribution analysis; (III) UCL Standing §17204; (IV) UCL Business Act §17200; (V) Defamation/Opinion — nonactionable opinion (Milkovich framework); (VI) Libel Per Quod/Special Damages §45a."),
    ("Case Citations — 10/10",
     "All 6 required statutes cited in body: 17 U.S.C. §107, §504(b), Cal. Civ. Code §45, §45a, Cal. Bus. & Prof. Code §17200, §17204. All 5 key cases cited: Campbell v. Acuff-Rose, Harper & Row, Milkovich v. Lorain Journal, Kasky v. Nike, Kwikset v. Superior Court."),
    ("Analytical Quality — 27/30",
     "Very strong analysis. All key markers present: specific case facts tied to each argument (blog, no advertisers, 'the heart,' ruthless criticism, little content beyond quotation); counterarguments acknowledged and rebutted for all four fair-use factors; the critical distinction between harm from criticism and harm from market substitution explicitly and repeatedly drawn; voluntary contributions correctly analyzed as non-commercial and attribution-problematic under §504(b). Slight deduction: the defamation section (Sec. V) focuses heavily on the opinion/Milkovich framework but does not fully develop the 'accurate quotation cannot supply falsity' sub-argument from golden evaluation §5B.1."),
]
for title, body in sub_items:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(title + ": ")
    run.bold = True
    p.add_run(body)

doc.add_paragraph()

# ── Deviations ────────────────────────────────────────────────────────────────
heading("Deviations from Golden Evaluation", 1)
deviations = [
    "17 U.S.C. §107 and Cal. Bus. & Prof. Code §17204 are cited extensively in the brief body but are absent from the auto-generated Table of Authorities index. This appears to be a TOA field-code indexing gap, not a missing citation. Both statutes appear correctly in the argument sections.",
    "Defamation section (Sec. V) does not explicitly argue that 'accurate quotation cannot supply falsity' (golden evaluation §5B.1). The section focuses on the opinion/disclosed-factual-basis doctrine (Milkovich) and rhetorical hyperbole, which covers §5B.2–3 but leaves §5B.1 underdeveloped.",
]
for d in deviations:
    doc.add_paragraph(d, style="List Bullet")

doc.add_paragraph()

# ── New Elements ──────────────────────────────────────────────────────────────
heading("New Elements Not in Golden Evaluation", 1)
para("Cases cited in this run that do not appear anywhere in the golden evaluation (not even as 'may appear'). Verification links provided for SME review.", italic=True)
doc.add_paragraph()

new_tbl = doc.add_table(rows=1, cols=4)
new_tbl.style = "Table Grid"
hdr = new_tbl.rows[0].cells
for i, h in enumerate(["Case / Statute", "Citation", "CourtListener Link", "Context in Brief"]):
    hdr[i].text = h
    for run in hdr[i].paragraphs[0].runs:
        run.bold = True
shade_row(new_tbl.rows[0], "D5E8F0")

new_elements = [
    ("Mackie v. Rieser",
     "296 F.3d 909 (9th Cir. 2002)",
     "https://www.courtlistener.com/?q=Mackie+v.+Rieser&type=o&order_by=score+desc",
     "Copyright damages — causation for actual damages under §504(b)"),
    ("Noelia Monge v. Maya Magazines, Inc.",
     "688 F.3d 1164 (9th Cir. 2012)",
     "https://www.courtlistener.com/?q=Monge+v.+Maya+Magazines&type=o&order_by=score+desc",
     "Fair use Factor 2 — published vs. unpublished distinction"),
    ("New York Times Co. v. Sullivan",
     "376 U.S. 254 (1964)",
     "https://www.courtlistener.com/?q=New+York+Times+v.+Sullivan&type=o&order_by=score+desc",
     "Defamation section — constitutional limits on defamation liability"),
    ("Dr. Seuss Enterprises, L.P. v. Comicmix LLC",
     "983 F.3d 443 (9th Cir. 2020)",
     "https://www.courtlistener.com/?q=Dr.+Seuss+v.+Comicmix&type=o&order_by=score+desc",
     "Defamation/opinion section — opinion and rhetorical hyperbole"),
    ("Cel-Tech Communications, Inc. v. Los Angeles Cellular Telephone Co.",
     "20 Cal. 4th 163 (1999)",
     "https://www.courtlistener.com/?q=Cel-Tech+Communications+v.+Los+Angeles+Cellular&type=o&order_by=score+desc",
     "UCL section — UCL's 'unfair' prong standard"),
    ("Jenkins v. JPMorgan Chase Bank, N.A.",
     "216 Cal. App. 4th 497 (2013)",
     "https://www.courtlistener.com/?q=Jenkins+v.+JPMorgan+Chase&type=o&order_by=score+desc",
     "UCL Standing — non-speculative economic injury requirement"),
    ("Barnes-Hind, Inc. v. Superior Court",
     "181 Cal. App. 3d 377 (1986)",
     "https://www.courtlistener.com/?q=Barnes-Hind+v.+Superior+Court&type=o&order_by=score+desc",
     "Libel per quod — §45a special damages requirement"),
    ("Gregory v. McDonnell Douglas Corp.",
     "17 Cal. 3d 596 (1976)",
     "https://www.courtlistener.com/?q=Gregory+v.+McDonnell+Douglas&type=o&order_by=score+desc",
     "Statement of Issues — opinion vs. fact framework"),
    ("Gunn v. Minton",
     "568 U.S. 251 (2013)",
     "https://www.courtlistener.com/?q=Gunn+v.+Minton&type=o&order_by=score+desc",
     "Jurisdiction — 'arising under' analysis for embedded federal issues"),
    ("United Mine Workers v. Gibbs",
     "383 U.S. 715 (1966)",
     "https://www.courtlistener.com/?q=United+Mine+Workers+v.+Gibbs&type=o&order_by=score+desc",
     "Jurisdiction — pendent jurisdiction over state-law claims"),
    ("Calder v. Jones",
     "465 U.S. 783 (1984)",
     "https://www.courtlistener.com/?q=Calder+v.+Jones+465+U.S.&type=o&order_by=score+desc",
     "Jurisdiction — specific personal jurisdiction, effects test"),
]
for case, citation, link, context in new_elements:
    r = new_tbl.add_row()
    r.cells[0].text = case
    r.cells[1].text = citation
    # Add hyperlink text (plain link since we can't easily add hyperlinks in table cells via python-docx)
    r.cells[2].text = link
    r.cells[3].text = context

doc.add_paragraph()
para("Note: Cases already in the golden evaluation as 'may appear' that also appeared in this run (Polar Bear Prods., Seltzer v. Green Day, Baker v. L.A. Herald Examiner, Slaughter v. Friedman, Pavlovich, Vons, Daimler, Korea Supply, In re Tobacco II, Perfect 10 v. Amazon) are not listed above.", italic=True)

# Save
out = "C:/Users/nicho/OneDrive/Nick Work Laptop/Claude-SME-QA-Bot--master/solutions/legal-research/evals/jane-smith/runs/run-5/qa-scorecard.docx"
doc.save(out)
print(f"Saved: {out}")
